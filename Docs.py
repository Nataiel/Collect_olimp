import json
from docx import Document
import os


def create_statements_simple_and_reliable():
    """
    Самый простой и надежный способ создания заявлений.
    """
    # Загружаем данные
    with open("olympiad_data — копия.json", 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Создаем новый документ
    doc = Document()

    # Открываем шаблон один раз, чтобы получить его содержимое
    template_doc = Document("Заявление ВСоШ.docx")

    # Получаем все параграфы из шаблона как список текстов
    template_paragraphs = []
    for para in template_doc.paragraphs:
        template_paragraphs.append(para.text)

    statement_count = 0

    # Для каждого ученика создаем заявление
    for class_name, students in data.items():

        for student_name, subjects in students.items():
            # Получаем выбранные предметы
            selected = [subj for subj, chosen in subjects.items() if chosen]

            if not selected:  # Пропускаем если нет выбранных предметов
                continue

            statement_count += 1
            subjects_str = ", ".join(selected)

            print(f"Создаю заявление {statement_count}: {student_name}")

            # Добавляем каждый параграф из шаблона с заменой
            for para_text in template_paragraphs:
                # Заменяем плейсхолдеры
                new_text = para_text
                new_text = new_text.replace('%ФИО%', student_name)
                new_text = new_text.replace('%Класс%', class_name)
                new_text = new_text.replace('%олимпиады%', subjects_str)
                new_text = new_text.replace('ОО', 'ОО')

                # Добавляем параграф в документ
                if new_text.strip():  # Не добавляем пустые строки
                    doc.add_paragraph(new_text)

                doc.add_page_break()

    # Сохраняем документ
    output_file = "Заявления_готовые.docx"
    doc.save(output_file)

    print(f"\n✅ Готово! Создано {statement_count} заявлений")
    print(f"📄 Файл сохранен: {output_file}")


# Запускаем программу
if __name__ == "__main__":
    # Проверяем файлы
    if not os.path.exists("Заявление ВСоШ.docx"):
        print("❌ Файл 'Заявление ВСоШ.docx' не найден!")
    elif not os.path.exists("olympiad_data — копия.json"):
        print("❌ Файл 'olympiad_data — копия.json' не найден!")
    else:
        try:
            create_statements_simple_and_reliable()
        except Exception as e:
            print(f"❌ Ошибка: {e}")
            print("Убедитесь, что установлена библиотека python-docx:")
            print("pip install python-docx")
