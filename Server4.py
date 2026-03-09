import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading
import pandas as pd
import json
import os
from flask import Flask, request, jsonify, render_template_string
import webbrowser
import re
from docx import Document
from docx.enum.text import WD_BREAK
from collections import defaultdict
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from openpyxl.styles import PatternFill, Border, Side
from openpyxl.utils import get_column_letter
import socket


# Глобальные переменные
server_thread = None
server_running = False
data_file = 'olympiad_data.json'  # Файл хранения выбора олимпиад для каждого участника
excel_file = None
school_data = {}
template_path = None  # Путь к фалу шаблона заявления
template_diploma_path = None  # Путь к фалу шаблона грамоты
progress_bar = None  # Прогресс бар для отслеживания создания заявлений
progress_bar_diploma = None  # Прогресс бар для отслеживания создания грамот
data_diploma = None  # Данные для формирования грамот ВсОШ ШЭ

# Предметы для разных классов
subjects_4_class = ['Математика', 'Русский язык']
subjects_5_11_class = [
    'Англ. язык', 'Астрономия', 'Биология', 'География',
    'Инф.без.', 'ИИ', 'Искусство', 'Исп.язык', 'История', 'Итал.язык', 'Кит.язык',
    'Литература', 'Математика', 'Нем.язык', 'ОБЗР', 'Обществознание',
    'Право', 'Программирование', 'Робототехника', 'Русский язык', 'Труд',
    'Физика', 'Физ.культура', 'Фран.язык', 'Химия', 'Экология', 'Экономика'
]

olympiad_data = {}
first_run = True
flask_app = None


def get_subjects_for_class(class_name):
    """Определяет набор предметов в зависимости от класса"""
    match = re.search(r'(\d+)', class_name)
    if match:
        grade = int(match.group(1))
        if grade == 4:
            return subjects_4_class
        else:
            return subjects_5_11_class
    return subjects_5_11_class


def setup_flask_app():
    """Настройка Flask приложения"""
    global flask_app
    flask_app = Flask(__name__)

    MAIN_PAGE = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Участие в олимпиадах</title>
        <meta charset="utf-8">
        <style>
            body { 
                font-family: Arial, sans-serif; 
                max-width: 900px; 
                margin: 50px auto; 
                padding: 20px;
                background: #f5f5f5;
            }
            .header {
                text-align: center;
                background: white;
                padding: 30px;
                border-radius: 10px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                margin-bottom: 30px;
            }
            .grade-section {
                background: white;
                padding: 25px;
                border-radius: 10px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                margin-bottom: 20px;
                display: flex;
                align-items: center;
                gap: 25px;
            }
            .grade-title {
                font-size: 24px;
                font-weight: bold;
                color: #2c3e50;
                min-width: 80px;
                text-align: center;
            }
            .class-grid {
                display: flex;
                gap: 12px;
                flex-wrap: wrap;
                flex: 1;
            }
            .class-card {
                background: #3498db;
                color: white;
                padding: 18px 22px;
                text-align: center;
                border-radius: 8px;
                text-decoration: none;
                font-size: 18px;
                font-weight: bold;
                box-shadow: 0 2px 8px rgba(0,0,0,0.2);
                transition: transform 0.2s, box-shadow 0.2s, background 0.2s;
                min-width: 50px;
            }
            .class-card:hover {
                transform: translateY(-3px);
                box-shadow: 0 5px 15px rgba(0,0,0,0.3);
                background: #2980b9;
            }
            .stats-link {
                display: block;
                margin-top: 30px;
                text-align: center;
                color: #3498db;
                text-decoration: none;
                font-weight: bold;
                font-size: 18px;
                padding: 15px;
                background: white;
                border-radius: 8px;
                box-shadow: 0 2px 8px rgba(0,0,0,0.1);
                transition: all 0.2s;
            }
            .stats-link:hover {
                text-decoration: none;
                background: #3498db;
                color: white;
                box-shadow: 0 4px 12px rgba(52,152,219,0.3);
            }
            .no-classes {
                text-align: center;
                color: #7f8c8d;
                font-style: italic;
                padding: 20px;
                background: white;
                border-radius: 10px;
            }
        </style>
    </head>
    <body>
        <div class="header">
            <h1>🏆 Участие в олимпиадах</h1>
            <div>Выберите класс для редактирования</div>
        </div>

        {% if classes_by_grade %}
            {% for grade, class_letters in classes_by_grade.items() %}
            <div class="grade-section">
                <div class="grade-title">{{ grade }} класс</div>
                <div class="class-grid">
                    {% for class_letter in class_letters %}
                    <a href="/class/{{ grade }}{{ class_letter }}" class="class-card">{{ class_letter }}</a>
                    {% endfor %}
                </div>
            </div>
            {% endfor %}
        {% else %}
            <div class="no-classes">Классы не загружены</div>
        {% endif %}

        <a href="/stats" class="stats-link">📊 Посмотреть статистику</a>
    </body>
    </html>
    """

    CLASS_PAGE_TEMPLATE = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Класс {{ class_name }} - Олимпиады</title>
        <meta charset="utf-8">
        <style>
            body { 
                font-family: Arial, sans-serif; 
                max-width: 95vw; 
                margin: 10px auto; 
                padding: 10px;
                background: #f5f5f5;
            }
            .header {
                background: white;
                padding: 12px;
                border-radius: 10px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                margin-bottom: 10px;
                position: sticky;
                top: 0;
                z-index: 100;
            }
            .back-link {
                display: inline-block;
                text-decoration: none;
                color: #3498db;
                font-weight: bold;
            }
            .back-link:hover {
                text-decoration: underline;
            }
            .table-container {
                max-height: 70vh;
                overflow: auto;
                background: white;
                border-radius: 10px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }
            table {
                width: 100%;
                border-collapse: collapse;
                background: white;
            }
            thead {
                position: sticky;
                top: 0;
                z-index: 10;
            }
            th, td {
                padding: 6px 4px;
                text-align: center;
                border-bottom: 1px solid #ecf0f1;
                border-right: 1px solid #ecf0f1;
            }
            th {
                background: white;
                color: #34495e;
                position: relative;
                min-width: 50px;
                height: 150px;
            }
            .student-name {
                width: 15%;
                background: #f8f9fa;
                text-align: left;
                padding-left: 12px;
            }
            .subject-header {
                vertical-align: center;
                writing-mode: vertical-lr;
                transform: rotate(180deg);
                text-align: left;
                padding: 6px 3px;
                font-size: 14px;
                min-width: 1px;
                min-height: 140px;
            }
            .checkbox-cell {
                text-align: center;
                width: 50px;
            }
            input[type="checkbox"] {
                transform: scale(1.3);
                cursor: pointer;
            }
            tr:hover {
                background: #f8f9fa;
            }
            tr:hover .student-name {
                background: #e8f4fd;
            }
            .save-btn {
                background: #27ae60;
                color: white;
                padding: 12px 30px;
                border: none;
                border-radius: 5px;
                cursor: pointer;
                font-size: 14px;
                margin-top: 20px;
                margin-right: 10px;
            }
            .save-btn:hover {
                background: #219a52;
            }
            .reset-btn {
                background: #e74c3c;
                color: white;
                padding: 12px 30px;
                border: none;
                border-radius: 5px;
                cursor: pointer;
                font-size: 14px;
                margin-top: 20px;
            }
            .reset-btn:hover {
                background: #c0392b;
            }
            .button-group {
                text-align: center;
                position: sticky;
                bottom: 0;
                background: #f5f5f5;
                padding: 15px 0;
                margin-top: 10px;
            }
            .message {
                padding: 10px;
                margin: 10px 0;
                border-radius: 5px;
                text-align: center;
                font-weight: bold;
            }
            .success {
                background: #d4edda;
                color: #155724;
                border: 1px solid #c3e6cb;
            }
            .info {
                background: #d1ecf1;
                color: #0c5460;
                border: 1px solid #bee5eb;
            }
            .student-row:hover {
                background: #f0f8ff;
            }
            .table-container::-webkit-scrollbar {
                width: 8px;
                height: 8px;
            }
            .table-container::-webkit-scrollbar-track {
                background: #f1f1f1;
                border-radius: 4px;
            }
            .table-container::-webkit-scrollbar-thumb {
                background: #c1c1c1;
                border-radius: 4px;
            }
            .table-container::-webkit-scrollbar-thumb:hover {
                background: #a8a8a8;
            }
        </style>
    </head>
    <body>
        <div class="header">
            <a href="/" class="back-link">← Назад к выбору класса</a>
            <h1>🎓 {{ class_name }} класс - Участие в олимпиадах</h1>
        </div>

        {% if message %}
        <div class="message success">{{ message }}</div>
        {% endif %}

        <form action="/save/{{ class_name }}" method="post">
            <div class="table-container">
                <table>
                    <thead>
                        <tr>
                            <th class="student-name">Ученик</th>
                            {% for subject in subjects %}
                            <th class="subject-header">{{ subject }}</th>
                            {% endfor %}
                        </tr>
                    </thead>
                    <tbody>
                        {% for student in students %}
                        <tr class="student-row">
                            <td class="student-name">{{ student }}</td>
                            {% for subject in subjects %}
                            <td class="checkbox-cell">
                                <input type="checkbox" 
                                       name="{{ student }}_{{ subject }}" 
                                       {{ 'checked' if participation[student][subject] else '' }}>
                            </td>
                            {% endfor %}
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </div>

            <div class="button-group">
                <button type="submit" class="save-btn">💾 Сохранить изменения</button>
                <button type="button" class="reset-btn" onclick="resetClass()">🔄 Сбросить класс</button>
            </div>
        </form>

        <script>
            function resetClass() {
                if (confirm('Вы уверены, что хотите сбросить ВСЕ выборы для этого класса? Это действие нельзя отменить.')) {
                    window.location.href = '/reset/{{ class_name }}';
                }
            }

            document.addEventListener('DOMContentLoaded', function() {
                const checkboxes = document.querySelectorAll('input[type="checkbox"]');

                function updateSelectionCount() {
                    const checked = document.querySelectorAll('input[type="checkbox"]:checked').length;
                    console.log(`Выбрано: ${checked} предметов`);
                }

                checkboxes.forEach(cb => cb.addEventListener('change', updateSelectionCount));
            });
        </script>
    </body>
    </html>
    """

    STATS_PAGE = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Статистика олимпиад</title>
        <meta charset="utf-8">
        <style>
            body { 
                font-family: Arial, sans-serif; 
                max-width: 1200px; 
                margin: 50px auto; 
                padding: 20px;
                background: #f5f5f5;
            }
            .back-link { 
                display: inline-block; 
                margin-bottom: 20px; 
                text-decoration: none; 
                color: #3498db; 
                font-weight: bold; 
                padding: 10px 20px;
                background: white;
                border-radius: 5px;
                box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            }
            .back-link:hover {
                background: #3498db;
                color: white;
            }
            .stat-card { 
                background: white; 
                padding: 25px; 
                margin: 20px 0; 
                border-radius: 10px; 
                box-shadow: 0 2px 10px rgba(0,0,0,0.1); 
            }
            .class-title {
                color: #2c3e50;
                border-bottom: 2px solid #3498db;
                padding-bottom: 10px;
                margin-bottom: 15px;
            }
            .subjects-grid {
                display: grid;
                grid-template-columns: repeat(auto-fill, minmax(300px, 1fr));
                gap: 15px;
                margin-top: 15px;
            }
            .subject-item {
                background: #f8f9fa;
                padding: 15px;
                border-radius: 8px;
                border-left: 4px solid #3498db;
            }
            .subject-name {
                font-weight: bold;
                color: #2c3e50;
                margin-bottom: 5px;
            }
            .subject-stats {
                display: flex;
                justify-content: space-between;
                color: #666;
            }
            .subject-count {
                font-weight: bold;
                color: #27ae60;
            }
            .subject-percent {
                font-weight: bold;
                color: #e74c3c;
            }
            .total-stats {
                background: white;
                color: #34495e;
                padding: 15px;
                border-radius: 8px;
                margin-top: 15px;
                text-align: center;
            }
            .grade-section {
                margin-bottom: 30px;
            }
            .grade-title {
                font-size: 24px;
                color: #2c3e50;
                margin-bottom: 15px;
                padding: 10px;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                border-radius: 8px;
                text-align: center;
            }
        </style>
    </head>
    <body>
        <a href="/" class="back-link">← Назад к выбору класса</a>
        <h1>📊 Статистика участия в олимпиадах</h1>

        {% for grade, classes_data in stats_by_grade.items() %}
        <div class="grade-section">
            <div class="grade-title">{{ grade }} классы</div>
            {% for class_data in classes_data %}
            <div class="stat-card">
                <h2 class="class-title">🎓 Класс {{ class_data.class_name }}</h2>

                <div class="total-stats">
                    Всего учеников: <strong>{{ class_data.total_students }}</strong> | 
                    Всего выборов: <strong>{{ class_data.total_choices }}</strong> 
                </div>

                <div class="subjects-grid">
                    {% for subject_data in class_data.subjects %}
                    <div class="subject-item">
                        <div class="subject-name">{{ subject_data.name }}</div>
                        <div class="subject-stats">
                            <span class="subject-count">{{ subject_data.count }}/{{ class_data.total_students }}</span>
                            <span class="subject-percent">{{ subject_data.percent }}%</span>
                        </div>
                    </div>
                    {% endfor %}
                </div>
            </div>
            {% endfor %}
        </div>
        {% endfor %}
    </body>
    </html>
    """

    @flask_app.route('/')
    def home():
        if not school_data:
            return "Данные не загружены. Загрузите Excel файл через GUI."

        # Группируем классы по параллелям
        classes_by_grade = {}

        for class_name in school_data.keys():
            match = re.search(r'(\d+)\s*([А-ЯA-Z])', class_name)
            if match:
                grade = match.group(1)
                class_letter = match.group(2)

                if grade not in classes_by_grade:
                    classes_by_grade[grade] = []

                if class_letter not in classes_by_grade[grade]:
                    classes_by_grade[grade].append(class_letter)

        sorted_grades = sorted(classes_by_grade.keys(), key=lambda x: int(x))
        for grade in sorted_grades:
            classes_by_grade[grade] = sorted(classes_by_grade[grade])

        return render_template_string(MAIN_PAGE, classes_by_grade=classes_by_grade)

    @flask_app.route('/class/<class_name>')
    def class_page(class_name):
        if class_name not in school_data:
            return "Класс не найден", 404

        students = school_data[class_name]
        participation = olympiad_data.get(class_name, {})
        class_subjects = get_subjects_for_class(class_name)

        return render_template_string(
            CLASS_PAGE_TEMPLATE,
            class_name=class_name,
            students=students,
            subjects=class_subjects,
            participation=participation,
            message=request.args.get('message')
        )

    @flask_app.route('/save/<class_name>', methods=['POST'])
    def save_data(class_name):
        try:
            if class_name not in school_data:
                return "Класс не найден", 404

            class_subjects = get_subjects_for_class(class_name)

            for student in school_data[class_name]:
                for subject in class_subjects:
                    checkbox_name = f"{student}_{subject}"
                    olympiad_data[class_name][student][subject] = checkbox_name in request.form

            save_olympiad_data()

            return render_template_string(
                CLASS_PAGE_TEMPLATE,
                class_name=class_name,
                students=school_data[class_name],
                subjects=class_subjects,
                participation=olympiad_data[class_name],
                message='✅ Данные успешно сохранены!'
            )

        except Exception as e:
            return f"Ошибка: {e}", 500

    @flask_app.route('/reset/<class_name>')
    def reset_class(class_name):
        """Сброс данных для конкретного класса"""
        if class_name in olympiad_data:
            class_subjects = get_subjects_for_class(class_name)
            for student in olympiad_data[class_name]:
                for subject in class_subjects:
                    olympiad_data[class_name][student][subject] = False
            save_olympiad_data()

        class_subjects = get_subjects_for_class(class_name)

        return render_template_string(
            CLASS_PAGE_TEMPLATE,
            class_name=class_name,
            students=school_data[class_name],
            subjects=class_subjects,
            participation=olympiad_data[class_name],
            message='✅ Данные класса сброшены!'
        )

    @flask_app.route('/stats')
    def stats_page():
        """Страница статистики"""
        stats_by_grade = {}

        for class_name in school_data:
            # Определяем параллель
            match = re.search(r'(\d+)', class_name)
            if match:
                grade = f"{match.group(1)}"
            else:
                grade = "Другие"

            if grade not in stats_by_grade:
                stats_by_grade[grade] = []

            class_subjects = get_subjects_for_class(class_name)
            total_students = len(school_data[class_name])
            total_choices = 0

            subjects_data = []
            for subject in class_subjects:
                count = sum(1 for student_data in olympiad_data.get(class_name, {}).values()
                            if student_data.get(subject, False))
                percent = round((count / total_students * 100) if total_students > 0 else 0, 1)
                total_choices += count

                subjects_data.append({
                    'name': subject,
                    'count': count,
                    'percent': percent
                })

            # Сортируем предметы по убыванию популярности
            subjects_data.sort(key=lambda x: x['count'], reverse=True)

            avg_per_student = round(total_choices / total_students, 1) if total_students > 0 else 0

            stats_by_grade[grade].append({
                'class_name': class_name,
                'total_students': total_students,
                'total_choices': total_choices,
                'avg_per_student': avg_per_student,
                'subjects': subjects_data
            })

        # Сортируем классы внутри параллелей
        for grade in stats_by_grade:
            stats_by_grade[grade].sort(key=lambda x: x['class_name'])

        # Сортируем параллели по номеру
        sorted_stats = {}
        for grade in sorted(stats_by_grade.keys(), key=lambda x: int(x) if x.isdigit() else 999):
            sorted_stats[grade] = stats_by_grade[grade]

        return render_template_string(STATS_PAGE, stats_by_grade=sorted_stats)

    @flask_app.route('/api/stats')
    def api_stats():
        stats = {}
        for class_name in school_data:
            class_subjects = get_subjects_for_class(class_name)
            stats[class_name] = {
                'total_students': len(school_data[class_name]),
                'participation_by_subject': {}
            }
            for subject in class_subjects:
                count = sum(1 for student_data in olympiad_data.get(class_name, {}).values()
                            if student_data.get(subject, False))
                stats[class_name]['participation_by_subject'][subject] = count
        return jsonify(stats)


def load_students_from_excel(file_path):
    """Загрузка данных учеников из Excel файла"""
    global school_data
    try:
        school_data = {}
        excel_data = pd.read_excel(file_path, sheet_name=None)

        # Список разрешенных классов (4-11)
        allowed_classes = []
        for grade in range(4, 12):  # 4-11 классы
            for letter in ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С']:
                allowed_classes.append(f"{grade}{letter}")

        for sheet_name, df in excel_data.items():
            normalized_name = ''.join(sheet_name.split())
            sheet_name = ''.join(sheet_name.split())

            match = re.match(r'(\d+)\s*([А-ЯA-Z])', normalized_name)
            if not match:
                continue

            grade = int(match.group(1))
            if grade < 4 or grade > 11:
                continue

            if df.empty:
                continue

            students = []
            for index, row in df.iterrows():
                first_cell = str(row.iloc[0]).strip().lower()
                if first_cell in ['фамилия', 'фио', 'ученик', 'фамилия имя отчество']:
                    continue

                last_name = str(row.iloc[0]).strip() if pd.notna(row.iloc[0]) else ''
                first_name = str(row.iloc[1]).strip() if len(row) > 1 and pd.notna(row.iloc[1]) else ''
                middle_name = str(row.iloc[2]).strip() if len(row) > 2 and pd.notna(row.iloc[2]) else ''

                if not last_name or last_name == 'nan':
                    continue

                fio_parts = [last_name]
                if first_name and first_name != 'nan':
                    fio_parts.append(first_name)
                if middle_name and middle_name != 'nan':
                    fio_parts.append(middle_name)

                full_name = ' '.join(fio_parts)
                students.append(full_name)

            if students:
                school_data[sheet_name] = students
                print(f"Загружен класс {sheet_name}: {len(students)} учеников")

        print(f"Всего загружено классов: {len(school_data)}")
        print(f"Загруженные классы: {list(school_data.keys())}")
        return True

    except Exception as e:
        print(f"Ошибка загрузки Excel: {e}")
        return False


def reset_olympiad_data():
    """Сброс данных об участии (все чекбоксы не отмечены)"""
    global olympiad_data
    olympiad_data = {}
    for class_name in school_data:
        olympiad_data[class_name] = {}
        class_subjects = get_subjects_for_class(class_name)
        for student in school_data[class_name]:
            olympiad_data[class_name][student] = {subject: False for subject in class_subjects}

    save_olympiad_data()
    print("✅ Данные олимпиад сброшены (все чекбоксы не отмечены)")


def load_olympiad_data():
    """Загрузка данных об участии"""
    global olympiad_data
    if os.path.exists(data_file):
        print("📁 Загружаем существующие данные из файла...")
        with open(data_file, 'r', encoding='utf-8') as f:
            olympiad_data = json.load(f)

        data_updated = False
        for class_name in school_data:
            if class_name not in olympiad_data:
                olympiad_data[class_name] = {}
                data_updated = True

            class_subjects = get_subjects_for_class(class_name)

            for student in school_data[class_name]:
                if student not in olympiad_data[class_name]:
                    olympiad_data[class_name][student] = {subject: False for subject in class_subjects}
                    data_updated = True
                else:
                    for subject in class_subjects:
                        if subject not in olympiad_data[class_name][student]:
                            olympiad_data[class_name][student][subject] = False
                            data_updated = True
                    current_subjects = list(olympiad_data[class_name][student].keys())
                    for existing_subject in current_subjects:
                        if existing_subject not in class_subjects:
                            del olympiad_data[class_name][student][existing_subject]
                            data_updated = True

        if data_updated:
            save_olympiad_data()
            print("✅ Структура данных обновлена")
    else:
        print("📝 Файл данных не найден, создаем новый...")
        reset_olympiad_data()


def save_olympiad_data():
    """Сохранение данных об участии"""
    with open(data_file, 'w', encoding='utf-8') as f:
        json.dump(olympiad_data, f, ensure_ascii=False, indent=2)
    print("💾 Данные сохранены в файл")


def get_statistics():
    """Получение статистики для GUI"""
    stats = {}
    for class_name in school_data:
        stats[class_name] = {}
        class_subjects = get_subjects_for_class(class_name)
        for subject in class_subjects:
            count = sum(1 for student_data in olympiad_data.get(class_name, {}).values()
                        if student_data.get(subject, False))
            total = len(school_data[class_name])
            stats[class_name][subject] = f"{count}/{total}"
    return stats


def run_server():
    """Запуск сервера"""
    global server_running, first_run
    server_running = True

    if first_run and os.path.exists(data_file):
        print("📋 Файл данных существует - загружаем предыдущие выборы")
    elif first_run:
        print("🆕 Первый запуск - создаем чистые данные")

    first_run = False
    flask_app.run(host='0.0.0.0', port=5000, debug=False, use_reloader=False)


def stop_server():
    """Остановка сервера"""
    global server_running
    server_running = False


# GUI функции
def load_excel_file_gui(file_label, stats_tree):
    """Загрузка Excel файла с данными учеников"""
    global excel_file
    file_path = filedialog.askopenfilename(
        title="Выберите Excel файл с данными учеников",
        filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
    )

    if file_path:
        try:
            if load_students_from_excel(file_path):
                excel_file = file_path

                if not school_data:
                    messagebox.showwarning("Предупреждение",
                                           "Не найдено классов 4-11 в файле.\n\n"
                                           "Убедитесь, что:\n"
                                           "1. В файле есть листы для классов 4-11\n"
                                           "2. Данные содержат столбцы с Фамилией, Именем и Отчеством")
                    return

                if os.path.exists(data_file):
                    response = messagebox.askyesno(
                        "Загрузка данных",
                        "Обнаружен файл с предыдущими выборами.\n\nЗагрузить существующие данные? (Да)\nИли сбросить все выборы? (Нет)"
                    )
                    if response:
                        print("📁 Загружаем существующие данные...")
                        load_olympiad_data()
                    else:
                        print("🔄 Сбрасываем все данные...")
                        reset_olympiad_data()
                else:
                    print("🆕 Создаем новые данные...")
                    load_olympiad_data()

                file_label.config(text=f"Загружен: {os.path.basename(file_path)}")
                messagebox.showinfo("Успех", f"Данные успешно загружены!\nКлассы: {', '.join(school_data.keys())}")
                update_statistics_gui(stats_tree)
            else:
                messagebox.showerror("Ошибка", "Не удалось загрузить данные из файла")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при загрузке файла: {e}")


def reset_all_data_gui(stats_tree):
    """Сброс всех данных олимпиад"""
    if messagebox.askyesno("Сброс данных",
                           "Вы уверены, что хотите сбросить ВСЕ данные олимпиад?\nВсе выборы будут удалены!"):
        reset_olympiad_data()
        update_statistics_gui(stats_tree)
        messagebox.showinfo("Сброс", "Все данные олимпиад сброшены!")


def start_server_gui(status_label):
    """Запуск сервера"""
    global server_thread
    if not school_data:
        messagebox.showwarning("Предупреждение", "Сначала загрузите Excel файл с данными")
        return

    if not server_running:
        server_thread = threading.Thread(target=run_server)
        server_thread.daemon = True
        server_thread.start()
        hostname = socket.gethostname()
        local_ip = socket.gethostbyname(hostname)
        status_label.config(text=f"Статус: Сервер запущен на http://{local_ip}:5000",
                            foreground="green")
        messagebox.showinfo("Сервер запущен", "Сервер успешно запущен!\nВы можете открыть его в браузере.")
        create_internet_shortcut()


def create_internet_shortcut():
    """
    Создает ярлык интернет-ресурса (.url)
    """

    hostname = socket.gethostname()
    local_ip = socket.gethostbyname(hostname)

    # Создаем файл .url
    with open("Выбор предметов ВсОШ.url", 'w', encoding='utf-8') as f:
        f.write(f'[InternetShortcut]\n')
        f.write(f'URL=http://{local_ip}:5000\n')
        f.write(f'IDList=\n')
        f.write(f'IconIndex=0\n')
        f.write(f'HotKey=0\n')

    print(f"Ярлык создан: Выбор предметов ВсОШ.url")


def stop_server_gui(status_label):
    """Остановка сервера"""
    if server_running:
        stop_server()
        status_label.config(text="Статус: Сервер остановлен", foreground="red")
        messagebox.showinfo("Сервер остановлен", "Сервер был остановлен.")
    else:
        messagebox.showinfo("Информация", "Сервер уже остановлен")


def open_browser_gui():
    """Открытие браузера"""
    if server_running:
        hostname = socket.gethostname()
        local_ip = socket.gethostbyname(hostname)
        webbrowser.open(f"http://{local_ip}:5000")
    else:
        messagebox.showwarning("Предупреждение", "Сервер не запущен")


def update_statistics_gui(stats_tree):
    """Обновление статистики в таблице"""
    # Очищаем таблицу
    for item in stats_tree.get_children():
        stats_tree.delete(item)

    # Если нет данных, выходим
    if not school_data or not olympiad_data:
        return

    # Заполняем таблицу
    for class_name, students in school_data.items():
        # Подсчет количества учеников, участвующих хотя бы в одной олимпиаде
        participating_students = 0

        if class_name in olympiad_data:
            for student, subjects_data in olympiad_data[class_name].items():
                # Проверяем, участвует ли ученик хотя бы в одной олимпиаде
                if any(subjects_data.values()):
                    participating_students += 1

        # Расчет процента участия
        total_students = len(students)
        if total_students > 0:
            participation_percentage = (participating_students / total_students) * 100
        else:
            participation_percentage = 0

        # Форматирование процента с одной десятичной цифрой
        percentage_text = f"{participation_percentage:.1f}%"

        # Добавляем в таблицу
        stats_tree.insert("", "end", values=(class_name, percentage_text))


def save_diploma_to_xlsx():
    global data_diploma
    protocol = filedialog.askopenfilename(
        title="Выберите протокол с результатами ВсОШ ШЭ",
        filetypes=[("Excel documents", "*.xlsx"), ("All files", "*.*")])

    # Проверка, выбран ли файл
    if not protocol:
        messagebox.showwarning("Предупреждение", "Файл не выбран")
        return

    sheets = pd.read_excel(protocol, sheet_name=None)
    res = pd.concat(sheets.values(), ignore_index=True)
    res = res[res["Статус"].isin(['Победитель', 'Призёр'])]
    res['Результат'] = res['Предмет'] + ' (' + res['Статус'] + ')'
    res['Класс участника'] = res['Класс участника'].str.upper().str.replace(":", "", regex=False)
    res[["Код", "ОО"]] = res["Школа"].str.split(" - ", n=1, expand=True)
    res = res[["Код", "ОО", 'Класс участника', 'Участник', 'Результат']]
    res = res.groupby(["Код", "ОО", "Класс участника", "Участник"])["Результат"].agg(", ".join).reset_index()

    # Создаем словарь {ФИО: (Класс, Результат, ОО)}
    aux_data_diploma = zip(res['Класс участника'], res['Участник'], res['Результат'], res['ОО'])
    data_diploma = dict()
    for x in aux_data_diploma:
        if not x[0] in data_diploma:
            data_diploma[x[0]] = [x[1:]]
        else:
            data_diploma[x[0]].append(x[1:])

    # Формируем имя для нового файла
    directory = os.path.dirname(protocol)
    filename = os.path.basename(protocol)
    new_filename = f"Обработан {filename}"
    new_protocol = os.path.join(directory, new_filename)

    res.to_excel(new_protocol, index=False)
    messagebox.showinfo("Успех", f"Данные сохранены в файл:\n{new_protocol}")


def save_to_xlsx():
    global data_file

    # Загрузка JSON файла
    with open('olympiad_data.json', 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Группируем данные по предметам (ТОЛЬКО УЧАСТНИКИ)
    subject_data = defaultdict(list)
    students_data = 0  # Всего учеников

    # Собираем данные по каждому предмету (только участники)
    for class_name, students in data.items():
        for student_name, subjects in students.items():
            students_data += 1
            for subject, status in subjects.items():
                if status:  # Только если ученик участвует (status == True)
                    subject_data[subject].append({'Класс': class_name, 'Ученик': student_name})

    # Создаем новую книгу (старая будет полностью заменена)
    wb = Workbook()

    # Удаляем лист по умолчанию
    default_sheet = wb.active
    wb.remove(default_sheet)

    # Для каждого предмета создаем новый лист
    for subject, records in sorted(subject_data.items()):
        if not records:
            continue  # Пропускаем пустые предметы

        # Получаем имя листа (ограничиваем 31 символом - ограничение Excel)
        sheet_name = str(subject)[:31]

        # Создаем новый лист
        sheets = wb.sheetnames
        if not sheet_name in sheets:
            sheet = wb.create_sheet(title=sheet_name)
        else:
            continue

        # Определяем заголовки из ключей первого словаря
        headers = list(records[0].keys())

        # Записываем заголовки в первую строку
        for col_num, header in enumerate(headers, 1):
            sheet.cell(row=1, column=col_num).value = header
            # Настраиваем стиль для заголовков (опционально)
            sheet.cell(row=1, column=col_num).font = Font(bold=True)
            sheet.cell(row=1, column=col_num).alignment = Alignment(horizontal='center')
        sheet.column_dimensions['B'].width = 50

    sheets = wb.sheetnames
    class_data = list(data.keys())

    for sheet in sheets:
        dataset = []
        for aux in subject_data[sheet]:
            dataset.append((aux['Класс'], aux['Ученик']))
        dataset.sort(key=lambda x: int(x[0][0]))

        wb.active = wb[sheet]
        activ_sheet = wb.active
        for data in dataset:
            activ_sheet.append(data)

        # Заголовки таблицы статистики
        activ_sheet['D1'].value = 'Класс'
        activ_sheet['D1'].font = Font(bold=True)
        activ_sheet['D1'].alignment = Alignment(horizontal='center')
        activ_sheet['E1'].value = 'Кол-во'
        activ_sheet['E1'].font = Font(bold=True)
        activ_sheet['E1'].alignment = Alignment(horizontal='center')

        for row in range(len(class_data)):
            activ_sheet.cell(row=row + 2, column=4).value = class_data[row]
            activ_sheet.cell(row=row + 2, column=5).value = f'=COUNTIF(A:A, "{class_data[row]}")'

        activ_sheet.cell(row=row + 3, column=4).value = 'Всего: '
        activ_sheet.cell(row=row + 3, column=4).font = Font(bold=True)
        activ_sheet.cell(row=row + 3, column=4).alignment = Alignment(horizontal='center')
        end_cell = activ_sheet.cell(row=row + 2, column=5).coordinate
        activ_sheet.cell(row=row + 3, column=5).value = f'=SUM(E2:{end_cell})'

    # Страница статистики
    statistic_sheet = wb.create_sheet(title='Статистика')
    col = 2

    # Столбец с классами
    statistic_sheet['A1'].value = 'Класс'
    statistic_sheet['A1'].font = Font(bold=True)
    statistic_sheet['A1'].alignment = Alignment(horizontal='center')
    statistic_sheet.row_dimensions[1].height = 105  # Установка высоты строки, так как предметы пишем вертикально
    statistic_sheet.column_dimensions['A'].width = 11
    for row in range(len(class_data)):
        statistic_sheet.cell(row=row + 2, column=col - 1).value = class_data[row]
        statistic_sheet.cell(row=row + 2, column=col - 1).alignment = Alignment(horizontal='center')

    statistic_sheet.cell(row=row + 3, column=col - 1).value = 'Всего:'
    statistic_sheet.cell(row=row + 3, column=col - 1).font = Font(bold=True)
    statistic_sheet.cell(row=row + 3, column=col - 1).alignment = Alignment(horizontal='center')

    statistic_sheet.cell(row=row + 4, column=col - 1).value = '% участия:'
    statistic_sheet.cell(row=row + 4, column=col - 1).font = Font(bold=True)
    statistic_sheet.cell(row=row + 4, column=col - 1).alignment = Alignment(horizontal='center')

    # Копирование данных с листов
    for sheet in sheets:
        # Заголовок предмета
        start = statistic_sheet.cell(row=1, column=col).coordinate

        statistic_sheet[start] = sheet
        statistic_sheet[start].font = Font(bold=True)
        statistic_sheet[start].alignment = Alignment(horizontal='center', text_rotation=90)
        col_letter = get_column_letter(col)
        statistic_sheet.column_dimensions[col_letter].width = 7
        # Копирование данных с листов
        for row in range(len(class_data)):
            statistic_sheet.cell(row=row + 2, column=col).value = f"='{sheet}'!E{row + 2}"

        stop = statistic_sheet.cell(row=row + 2, column=col).coordinate
        statistic_sheet.cell(row=row + 3, column=col).value = f'=SUM({start}:{stop})'
        statistic_sheet.cell(row=row + 3, column=col).font = Font(bold=True)
        statistic_sheet.cell(row=row + 3, column=col).alignment = Alignment(horizontal='center')

        statistic_sheet.cell(row=row + 4, column=col).value = f'=SUM({start}:{stop})/{students_data}'
        statistic_sheet.cell(row=row + 4, column=col).alignment = Alignment(horizontal='center')
        statistic_sheet.cell(row=row + 4, column=col).number_format = '0.0%'

        # Оформление: зебра столбцов через предмет
        row_end = row + 4
        gray_fill = PatternFill(start_color="F3F3F3", end_color="F3F3F3", fill_type="solid")
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                             top=Side(style='thin'), bottom=Side(style='thin'))
        if col % 2 == 0:
            for row in statistic_sheet.iter_rows(min_row=1, max_row=row_end, min_col=col, max_col=col):
                for cell in row:
                    cell.fill = gray_fill

        # Применяем границы ко всем ячейкам в диапазоне
        for row in statistic_sheet.iter_rows(min_row=1, max_row=row_end, min_col=1, max_col=col):
            for cell in row:
                cell.border = thin_border

        col += 1

    wb.save('Отчёт о выборе предметов ВсОШ ШЭ.xlsx')
    print('Выгрузка в Excel завершена')


def create_diploma():
    global data_diploma
    x = 0

    if not template_diploma_path:
        messagebox.showerror("Ошибка", "Сначала выберите шаблон грамоты")
        return

    # СОЗДАЕМ ПАПКУ "Грамоты" В ТЕКУЩЕЙ ДИРЕКТОРИИ
    current_dir = os.getcwd()
    diplomas_folder = os.path.join(current_dir, "Грамоты")

    if not os.path.exists(diplomas_folder):
        os.makedirs(diplomas_folder)
        print(f"Создана папка: {diplomas_folder}")

    for class_name, data in data_diploma.items():
        # Создаем новый документ на основе шаблона
        current_doc = Document(template_diploma_path)

        # Удаляем содержимое (оставляем только настройки секции с фоном)
        last_paragraph = current_doc.paragraphs[-1]
        run = last_paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)

        # Теперь добавляем содержимое для каждого ученика
        for student_idx, (student_name, aux_subject, school) in enumerate(data):
            # Загружаем шаблон для получения содержимого
            template_doc = Document(template_diploma_path)

            status = ''
            if 'Победитель' in aux_subject and 'Призёр' in aux_subject:
                status = 'победитель и призёр'
            elif 'Призёр' in aux_subject:
                status = 'призёр'
            elif 'Победитель' in aux_subject:
                status = 'победитель'

            count_subject = 'у'
            if ',' in aux_subject:
                count_subject = 'ам'

            name_parts = student_name.split()
            replacements = {
                '%ОО%': school,
                '%Класс%': class_name,
                '%Предметы%': aux_subject,
                '%Фамилия%': name_parts[0] if len(name_parts) > 0 else '',
                '%Имя%': name_parts[1] if len(name_parts) > 1 else '',
                '%призёр и победитель%': status,
                '%окон%': count_subject,
            }

            replace_placeholders_in_document(template_doc, replacements)

            # Копируем содержимое в текущий документ
            for element in template_doc.element.body:
                current_doc.element.body.append(element)

            # Добавляем разрыв страницы после каждого ученика (кроме последнего)
            if student_idx < len(data) - 1:
                if current_doc.paragraphs:
                    last_paragraph = current_doc.paragraphs[-1]
                    run = last_paragraph.add_run()
                    run.add_break(WD_BREAK.PAGE)

        x += 100 / len(data_diploma)
        progress_bar_diploma['value'] = x
        progress_bar_diploma.update()

        filename = f'{class_name}_Грамоты_ВСоШ_ШЭ.docx'
        filename = os.path.join(diplomas_folder, filename)
        current_doc.save(filename)


def replace_placeholders_in_document(doc, replacements):
    """
    Заменяет плейсхолдеры во всём документе, сохраняя форматирование
    replacements: словарь {плейсхолдер: значение}
    """
    # Обрабатываем все параграфы
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            new_text = text

            # Заменяем все плейсхолдеры в этом run
            for placeholder, value in replacements.items():
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, str(value))

            # Обновляем текст run, если были изменения
            run.text = new_text

    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text
                        new_text = text
                        for placeholder, value in replacements.items():
                            if placeholder in new_text:
                                new_text = new_text.replace(placeholder, str(value))
                        run.text = new_text


def create_statement():
    global data_file
    x = 0

    # СОЗДАЕМ ПАПКУ "Заявления" В ТЕКУЩЕЙ ДИРЕКТОРИИ
    current_dir = os.getcwd()
    diplomas_folder = os.path.join(current_dir, "Заявления")

    # Создаем папку, если её нет
    if not os.path.exists(diplomas_folder):
        os.makedirs(diplomas_folder)
        print(f"Создана папка: {diplomas_folder}")

    # Загрузка JSON файла
    with open('olympiad_data.json', 'r', encoding='utf-8') as f:
        data = json.load(f)

    for class_name, students in data.items():
        # Загружаем шаблон документа
        current_doc = Document()
        flag_statement = False

        for student_name, subjects in students.items():
            aux_subject = []
            for subject, status in subjects.items():
                if status:  # Только если ученик участвует (status == True)
                    aux_subject.append(subject)

            if not aux_subject:
                continue

            flag_statement = True

            aux_subject = ', '.join(aux_subject)
            temp = (('Англ. язык', 'Английский язык'), ('Инф.без.', 'Информатика (информационная безопасность)'),
                    ('ИИ', 'Информатика (искусственный интеллект)'), ('Искусство', 'Искусство МХК'),
                    ('Исп.язык', 'Испанский язык'), ('Итал.язык', 'Итальянский язык'), ('Кит.язык', 'Китайский язык'),
                    ('Нем.язык', 'Немецкий язык'), ('ОБЗР', 'Основы безопасности и защиты Родины'),
                    ('Программирование', 'Информатика (программирование)'),
                    ('Робототехника', 'Информатика (робототехника)'), ('Физ.культура', 'Физическая культура'),
                    ('Фран.язык', 'Французский язык'))
            for a, b in temp:
                aux_subject = aux_subject.replace(a, b)

            if current_doc.paragraphs:
                last_paragraph = current_doc.paragraphs[-1]
                run = last_paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)

            # Создаем новый документ из шаблона
            new_doc = Document(template_path)

            # Копируем все элементы из нового документа в текущий
            for element in new_doc.element.body:
                current_doc.element.body.append(element)

            # Определяем, какое заявление обрабатываем
            paragraphs = current_doc.paragraphs
            start_index = 0

            # Находим начало текущего заявления
            for idx, para in enumerate(paragraphs):
                if "ЗАЯВЛЕНИЕ" in para.text and idx > start_index:
                    # Это начало нового заявления
                    start_index = idx
                    break

            # Обрабатываем абзацы текущего заявления
            for para in paragraphs[start_index:]:
                if '%ФИО%' in para.text:
                    # Заменяем текст в абзаце
                    for run in para.runs:
                        if '%ФИО%' in run.text:
                            run.text = run.text.replace('%ФИО%', student_name)
                        if '%Класс%' in run.text:
                            run.text = run.text.replace('%Класс%', class_name)
                        if '%олимпиады%' in run.text:
                            run.text = run.text.replace('%олимпиады%', aux_subject)

        x += 100 / len(data_file)
        progress_bar['value'] = x
        progress_bar.update()

        # Сохраняем
        filename = f'Заявление ВсОШ ШЭ {class_name}.docx'
        filename = os.path.join(diplomas_folder, filename)
        current_doc.save(filename)


def select_template():
    """Выбор шаблона заявления в формате *.docx"""
    global template_path

    file_path = filedialog.askopenfilename(
        title="Выберите шаблон заявления",
        filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
    )

    if file_path:
        template_path = file_path
        return True
    return False


def select_template_diploma():
    """Выбор шаблона грамот в формате *.docx"""
    global template_diploma_path

    file_path = filedialog.askopenfilename(
        title="Выберите шаблон грамоты",
        filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
    )

    if file_path:
        template_diploma_path = file_path
        return True
    return False


def update_template_status(template_status_label):
    """Обновление статуса шаблона"""
    global template_path

    if template_path:
        template_status_label.config(text=f"Шаблон: {template_path}", foreground="green")
    else:
        template_status_label.config(text="Шаблон не выбран", foreground="gray")


def update_template_diploma_status(template_status_label):
    """Обновление статуса шаблона"""
    global template_diploma_path

    if template_diploma_path:
        template_status_label.config(text=f"Шаблон: {template_diploma_path}", foreground="green")
    else:
        template_status_label.config(text="Шаблон не выбран", foreground="gray")


def select_template_with_status(template_status_label):
    """Выбор шаблона заявлений с обновлением статуса"""
    if select_template():
        update_template_status(template_status_label)


def select_template_diploma_with_status(template_status_label):
    """Выбор шаблона для грамот с обновлением статуса"""
    if select_template_diploma():
        update_template_diploma_status(template_status_label)


def setup_gui():
    """Создание графического интерфейса"""
    global progress_bar
    global progress_bar_diploma
    root = tk.Tk()
    root.title("Филиал ЕДУ Ленинского района. Сбор сведений ВсОШ Школьный этап v0.1a©")
    root.geometry("650x700")

    main_frame = ttk.Frame(root, padding="10")
    main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

    title_label = ttk.Label(main_frame, text="🏆 Управление сервером олимпиад", font=("Arial", 16, "bold"))
    title_label.grid(row=0, column=0, columnspan=4, pady=(0, 20))

    # Загрузка данных
    file_frame = ttk.LabelFrame(main_frame, text="Загрузка данных", padding="10")
    file_frame.grid(row=1, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=(0, 10))

    file_label = ttk.Label(file_frame, text="Файл не выбран")
    file_label.grid(row=0, column=0, sticky=(tk.W, tk.E))

    ttk.Button(file_frame, text="Выбрать Excel файл",
               command=lambda: load_excel_file_gui(file_label, stats_tree)).grid(row=0, column=1, padx=(10, 0))

    ttk.Button(file_frame, text="Сбросить все данные",
               command=lambda: reset_all_data_gui(stats_tree)).grid(row=0, column=2, padx=(10, 0))

    # Управление сервером (обновленный стиль)
    server_frame = ttk.LabelFrame(main_frame, text="Управление сервером", padding="10")
    server_frame.grid(row=2, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=(0, 10))

    status_label = ttk.Label(server_frame, text="Статус: Сервер остановлен", foreground="red",
                             font=("Arial", 10, "bold"))
    status_label.grid(row=1, column=0, columnspan=4, pady=(0, 0), sticky="w")

    # Кнопки управления сервером с одинаковой шириной и эмодзи
    ttk.Button(server_frame, text="▶ Запустить сервер", command=lambda: start_server_gui(status_label), width=30,
               style="Server.TButton").grid(row=0, column=0, padx=(0, 10), sticky="w")

    ttk.Button(server_frame, text="⏹ Остановить сервер", command=lambda: stop_server_gui(status_label), width=30,
               style="Server.TButton").grid(row=0, column=1, padx=(0, 10), sticky="w")

    ttk.Button(server_frame, text="🌐 Открыть в браузере", command=open_browser_gui, width=30,
               style="Server.TButton").grid(row=0, column=2, padx=(0, 10), sticky="w")

    # НОВЫЕ КНОПКИ: Экспорт и формирование заявлений
    export_frame = ttk.LabelFrame(main_frame, text="Заявления", padding="10")
    export_frame.grid(row=3, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=(0, 10))

    # Создаем кнопку выбора шаблона
    template_button = ttk.Button(export_frame, text="📄 Шаблон заявления", width=30)
    template_button.grid(row=0, column=0, padx=(0, 10))

    # Кнопка формирования заявлений
    ttk.Button(export_frame, text="📝 Сформировать заявления", command=create_statement,
               width=30).grid(row=0, column=1, padx=(0, 10))

    # Кнопка выгрузки в Excel
    ttk.Button(export_frame, text="📊 Выгрузить в Excel", command=save_to_xlsx,
               width=30).grid(row=0, column=2, padx=(0, 10))

    # Метка для отображения статуса шаблона
    template_status_label = ttk.Label(export_frame, text="Шаблон не выбран", foreground="gray")
    template_status_label.grid(row=1, column=0, columnspan=3, pady=(5, 0), sticky="w")

    # Привязываем команду к кнопке шаблона
    template_button.config(command=lambda: select_template_with_status(template_status_label))

    # Прогресс-бар (изначально скрыт)
    progress_bar = ttk.Progressbar(export_frame, mode='determinate', length=100)
    progress_bar.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 0))

    # ЕЩЁ НОВЫЕ КНОПКИ: Грамоты за школьный тур ВсОШ
    diploma_frame = ttk.LabelFrame(main_frame, text="Грамоты ВсОШ ШЭ", padding="10")
    diploma_frame.grid(row=4, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=(0, 10))

    # Создаем кнопку выбора шаблона для грамот
    template_diploma_button = ttk.Button(diploma_frame, text="📄 Шаблон грамоты", width=30)
    template_diploma_button.grid(row=0, column=0, padx=(0, 10))

    # Кнопка формирования заявлений
    ttk.Button(diploma_frame, text="📊 Протокол ВсОШ ШЭ", command=save_diploma_to_xlsx,
               width=30).grid(row=0, column=1, padx=(0, 10))

    # Кнопка выгрузки в Excel
    ttk.Button(diploma_frame, text="📝 Сформировать грамоты", command=create_diploma,
               width=30).grid(row=0, column=2, padx=(0, 10))

    # Метка для отображения статуса шаблона
    template_status_diploma_label = ttk.Label(diploma_frame, text="Шаблон не выбран", foreground="gray")
    template_status_diploma_label.grid(row=1, column=0, columnspan=3, pady=(5, 0), sticky="w")

    # Привязываем команду к кнопке шаблона
    template_diploma_button.config(command=lambda: select_template_diploma_with_status(template_status_diploma_label))

    # Прогресс-бар (изначально скрыт)
    progress_bar_diploma = ttk.Progressbar(diploma_frame, mode='determinate', length=80)
    progress_bar_diploma.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 0))

    # Статистика участия
    stats_frame = ttk.LabelFrame(main_frame, text="Статистика участия", padding="10")
    stats_frame.grid(row=5, column=0, columnspan=4, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))

    columns = ("Класс", "Процент участия")
    stats_tree = ttk.Treeview(stats_frame, columns=columns, show="headings", height=6)

    stats_tree.heading("Класс", text="Класс")
    stats_tree.column("Класс", width=200, anchor="center")

    stats_tree.heading("Процент участия", text="Процент участия")
    stats_tree.column("Процент участия", width=200, anchor="center")

    stats_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

    scrollbar = ttk.Scrollbar(stats_frame, orient="vertical", command=stats_tree.yview)
    scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
    stats_tree.configure(yscrollcommand=scrollbar.set)

    ttk.Button(stats_frame, text="Обновить статистику",
               command=lambda: update_statistics_gui(stats_tree)).grid(row=1, column=0, pady=(10, 0))

    # Настройка весов для растягивания
    root.columnconfigure(0, weight=1)
    root.rowconfigure(0, weight=1)
    main_frame.columnconfigure(0, weight=1)
    main_frame.rowconfigure(4, weight=1)
    file_frame.columnconfigure(0, weight=1)
    server_frame.columnconfigure(0, weight=1)

    export_frame.columnconfigure(0, weight=0)
    export_frame.columnconfigure(1, weight=0)
    export_frame.columnconfigure(2, weight=0)

    stats_frame.columnconfigure(0, weight=1)
    stats_frame.rowconfigure(0, weight=1)

    server_frame.columnconfigure(0, weight=0)
    server_frame.columnconfigure(1, weight=0)
    server_frame.columnconfigure(2, weight=0)
    server_frame.columnconfigure(3, weight=1)

    # Показываем окно с инструкцией после загрузки главного окна
    root.after(100, show_instruction_window)

    return root


def show_instruction_window():
    """Показывает окно с краткой инструкцией при запуске программы"""
    instruction_window = tk.Toplevel()
    instruction_window.title("Краткая инструкция")
    instruction_window.geometry("700x600")
    instruction_window.resizable(False, False)

    # Делаем окно модальным (поверх основного)
    instruction_window.transient()
    instruction_window.grab_set()

    # Основной фрейм
    main_frame = ttk.Frame(instruction_window, padding="15")
    main_frame.pack(fill=tk.BOTH, expand=True)

    # Заголовок
    title_label = ttk.Label(main_frame, text="📋 КРАТКАЯ ИНСТРУКЦИЯ", font=("Arial", 16, "bold"))
    title_label.pack(pady=(0, 15))

    # Создаем текстовое поле с прокруткой
    text_frame = ttk.Frame(main_frame)
    text_frame.pack(fill=tk.BOTH, expand=True)

    text_widget = tk.Text(text_frame, wrap=tk.WORD, font=("Courier New", 11), padx=10, pady=10, bg="#f8f9fa")
    scrollbar = ttk.Scrollbar(text_frame, command=text_widget.yview)
    text_widget.configure(yscrollcommand=scrollbar.set)

    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    # Текст инструкции
    instruction_text = """
🚀 ЗАГРУЗИТЕ ДАННЫЕ
- Нажмите "Выбрать Excel файл" — загрузите список учеников (листы с названиями классов: 4А, 5Б, 8В...). 
  *Используется выгрузка из ГИС СО "ЕЦП"  Электронный жкрнал. Файл формируется в разделе: Учет учащихся и педагогов → Классы → Сформировать списки: Для всех классов, Поле в отчёте - любое (обычно выбирают Пол)
- Нажмите "Шаблон заявления" — выберите файл шаблона (.docx)
    
🌐 СОБЕРИТЕ ДАННЫЕ
- Нажмите "Запустить сервер"
- Нажмите "Открыть в браузере" или раздайте ярлык на сервер "Выбор предметов ВсОШ"
- В браузере, для каждого ученика в классе отмечайте галочками, кто в каких олимпиадах участвует. 
  *Данные сохраняются на компьютере, где запущен сервер. Программу можно закрывать и открывать - данные о выборе сохраняются (при загрузке списков класса буде диалог для подтверждения использования уже собранных данных).
- Нажимайте "Сохранить изменения"
    
📄 ПОЛУЧИТЕ РЕЗУЛЬТАТ
- Нажмите "Сформировать заявления"** — формируются готовые заявления в папке "Заявления" по классам.
- Нажмите "Выгрузить в Excel" — получите статистику выбора предметов со списком учащихся, общая статистика выбора предметов и расчёт процента участия в классе и школе.
    
🏆 ФОРМИРОВАНИЕ ГРАМОТ ШКОЛЬНОГО ЭТАПА ВсОШ:
- Выберите "Шаблон грамоты"
- Загрузите "Протокол ВсОШ ШЭ" из РБДО:  Школьный этап → Результаты диагностик → Экспорт → (ШЭ)
- Нажмите "Сформировать грамоты" — формируются грамоты в папке "Грамоты" по классам
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Шаблоны заявления и грамот содержат поля подстановки для их автоматического заполнения из полученных данных. Текст шаблонов можно менять по своему усмотрению, помня только то, что вместо заменителей появится нужный текст подстановки.

Поле заявления         Описание
%ФИО%                  ФИО ученика
%Класс%                Класс ученика
%олимпиады%            Список предметов

Поле грамоты           Описание
%ОО%                   Образовательная организация (школа)
%Класс%                Класс ученика
%Предметы%             Предметы с результатами
%Фамилия%              Фамилия ученика
%Имя%                  Имя ученика
%призёр и победитель%  Статус (победитель, призёр или оба сразу)
%у/ам%                 Окончание слова "предмет"
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

    text_widget.insert("1.0", instruction_text)
    text_widget.configure(state="disabled")  # Делаем текст только для чтения

    # Кнопка закрытия
    button_frame = ttk.Frame(main_frame)
    button_frame.pack(fill=tk.X, pady=(15, 0))

    close_button = ttk.Button(button_frame, text="Понятно, закрыть", command=instruction_window.destroy, width=20)
    close_button.pack()

    # Центрируем окно относительно основного
    instruction_window.update_idletasks()
    x = (instruction_window.winfo_screenwidth() // 2) - (instruction_window.winfo_width() // 2)
    y = (instruction_window.winfo_screenheight() // 2) - (instruction_window.winfo_height() // 2)
    instruction_window.geometry(f'+{x}+{y}')

    # Не блокируем основное окно полностью (можно вернуться к нему)
    instruction_window.grab_release()

def main():
    """Главная функция"""
    print("🚀 Запуск приложения управления олимпиадами...")

    setup_flask_app()
    root = setup_gui()
    root.mainloop()


if __name__ == '__main__':
    main()
